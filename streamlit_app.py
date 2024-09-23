from itertools import zip_longest
import streamlit as st
import fitz  # PyMuPDF
import docx
import plotly.graph_objects as go
import pandas as pd
from streamlit_chat import message
from langchain.chat_models import ChatOpenAI
from langchain.schema import (
    SystemMessage,
    HumanMessage,
    AIMessage
)
import openai
import os

# Set Streamlit page configuration
st.set_page_config(page_title="⚖️ LawGPT ChatBot", page_icon="⚖️", layout="wide")
st.sidebar.title("OPENAI_API_KEY")

# Function to handle API key input
def handle_api_key():
    st.session_state.openapi_key = st.sidebar.text_input("Enter your OpenAI API key:", type="password")

# Check if the API key is already in session state
if "openapi_key" not in st.session_state:
    handle_api_key()

# Show radio buttons if API key is submitted
if "openapi_key" in st.session_state and st.session_state.openapi_key:
    option = st.sidebar.radio("Choose an option:", ("Option 1", "Option 2", "Option 3"))

# Initialize session state variables
if 'generated' not in st.session_state:
    st.session_state['generated'] = []  # Store AI generated responses

if 'past' not in st.session_state:
    st.session_state['past'] = []  # Store past user inputs

if 'entered_prompt' not in st.session_state:
    st.session_state['entered_prompt'] = ""  # Store the latest user input

# Initialize the ChatOpenAI model if API key is available
if "openapi_key" in st.session_state and st.session_state.openapi_key:
    chat = ChatOpenAI(
        temperature=0.5,
        model_name="gpt-3.5-turbo",
        openai_api_key=st.session_state.openapi_key,
        max_tokens=20
    )

# Set the main title
st.title("⚖️ LawGPT - Your Legal Assistant")



def build_message_list():
    """
    Build a list of messages including system, human and AI messages.
    """
    # Start zipped_messages with the SystemMessage
    zipped_messages = [SystemMessage(
        content = """Your name is LawGPT. ⚖️ You are an AI Legal Assistant, here to provide legal advice and information. Please provide accurate and helpful information, and always maintain a polite and professional tone.

                1. Greet the user politely, ask their name, and inquire how you can assist them with legal-related queries.
                2. Provide informative and relevant responses to questions about various legal topics such as contracts 📜, family law 👪, criminal law 🕵️, corporate law 🏢, and more.
                3. Avoid discussing sensitive, offensive, or harmful content. Refrain from engaging in any form of discrimination, harassment, or inappropriate behavior.
                4. If the user asks about a topic unrelated to law, politely steer the conversation back to legal matters or inform them that the topic is outside the scope of this conversation.
                5. Be patient and considerate when responding to user queries, and provide clear explanations.
                6. If the user expresses gratitude or indicates the end of the conversation, respond with a polite farewell. 👋
                7. Do not generate long paragraphs in response. Maximum Words should be 100.

                Remember, your primary goal is to assist and educate users in the field of law. Always prioritize their learning experience and well-being."""
    )]

    # Zip together the past and generated messages
    for human_msg, ai_msg in zip_longest(st.session_state['past'], st.session_state['generated']):
        if human_msg is not None:
            zipped_messages.append(HumanMessage(content=human_msg))  # Add user messages
        if ai_msg is not None:
            zipped_messages.append(AIMessage(content=ai_msg))  # Add AI messages

    return zipped_messages


def generate_response():
    """
    Generate AI response using the ChatOpenAI model.
    """
    # Build the list of messages
    zipped_messages = build_message_list()

    # Generate response using the chat model
    ai_response = chat(zipped_messages)

    return ai_response.content


# Define function to submit user input
def submit():
    # Set entered_prompt to the current value of prompt_input
    st.session_state.entered_prompt = st.session_state.prompt_input
    # Clear prompt_input
    st.session_state.prompt_input = ""


# Create a text input for user
st.text_input('YOU: 🗣️', key='prompt_input', on_change=submit)

if st.session_state.entered_prompt != "":
    # Get user query
    user_query = st.session_state.entered_prompt

    # Append user query to past queries
    st.session_state.past.append(user_query)

    # Generate response
    output = generate_response()

    # Append AI response to generated responses
    st.session_state.generated.append(output)


# Create a radio button menu for feature selection
selected_option = st.sidebar.radio("Select Feature", [
    "Chat", 
    "Legal Documents", 
    "Translation", 
    "Personalized Advice", 
    "Legal Experts", 
    "Compliance Checks", 
    "Knowledge Graphs", 
    "Drafting", 
    "Precedents"
])

if selected_option == "Chat":
    # Display the chat history
    if st.session_state['generated']:
        for i in range(len(st.session_state['generated'])-1, -1, -1):
            # Display AI response
            message(st.session_state["generated"][i], key=str(i), )
            # Display user message
            message(st.session_state['past'][i], is_user=True, key=str(i) + '_user', )
elif selected_option == "Legal Documents":
    st.header("Legal Document Analysis")
    
    uploaded_file = st.file_uploader("Upload Document", type=['pdf', 'docx'], label_visibility="collapsed")

    if uploaded_file is not None:
        # Read the content of the uploaded file
        text = ""
        if uploaded_file.type == "application/pdf":
            # Read PDF
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf_document:
                for page in pdf_document:
                    text += page.get_text()
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Read DOCX
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])

        # Display extracted text (for debugging purposes, can be removed in production)
        st.write("Extracted Text (First 2000 Characters):")
        st.write(text[:2000])  # Show first 2000 characters

        # Analyze text with OpenAI
        if text:
            try:
                response = openai.Completion.create(
                    engine="gpt-3.5-turbo",
                    prompt=(
                        "Please analyze the following legal document text. "
                        "Extract and highlight key points, headings, and provide a detailed summary:\n\n"
                        f"{text}"
                    ),
                    max_tokens=1500,  # Increased token limit for detailed analysis
                    temperature=0.5
                )
                st.header("Document Analysis")
                st.write(response.choices[0].text.strip())
            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")


elif selected_option == "Translation":
    st.header("Language Translation")
    text_to_translate = st.text_input("Enter text for translation", key="translation_input")
    target_language = st.selectbox("Select target language", ["English", "Spanish", "French", "German", "Chinese", "Japanese", "Korean", "Urdu"])
    
    # Mapping target languages to OpenAI prompt formats
    language_map = {
        "English": "en",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Chinese": "zh",
        "Japanese": "ja",
        "Korean": "ko",
        "Urdu": "ur"  # Added Urdu language code
    }
    
    if text_to_translate and target_language:
        # Translate text using OpenAI's chat completion endpoint
        try:
            translation_prompt = f"Translate the following text to {language_map[target_language]}:\n\n{text_to_translate}"
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": translation_prompt}
                ],
                max_tokens=1000,
                temperature=0.5,
                api_key=openapi_key
            )
            translated_text = response.choices[0].message['content'].strip()
            st.write("Translated Text:")
            st.write(translated_text)
        except Exception as e:
            st.error(f"An error occurred: {e}")


# Define the Personalized Legal Advice section
if selected_option == "Personalized Advice":
    st.header("Personalized Legal Advice 🕵️‍♂️")

    # User input fields for personalized advice
    with st.form(key='personalized_advice_form'):
        st.subheader("Enter your details for personalized advice:")
        name = st.text_input("Name")
        email = st.text_input("Email")
        legal_issue = st.text_area("Describe your legal issue", placeholder="Provide a detailed description of your legal issue...")
        additional_info = st.text_area("Additional information (optional)", placeholder="Any additional details you think might be relevant...")
        submit_button = st.form_submit_button("Get Legal Advice")

    if submit_button:
        if name and email and legal_issue:
            try:
                # Prepare the prompt for OpenAI API
                prompt = f"""
                You are a legal advisor providing personalized legal advice. Use the following information to generate a professional and helpful response:

                Name: {name}
                Email: {email}
                Legal Issue: {legal_issue}
                Additional Information: {additional_info if additional_info else "None"}

                Provide clear, relevant, and concise legal advice. Ensure your response is professional and easy to understand.
                """

                # Call OpenAI API for generating legal advice
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a legal advisor providing personalized legal advice."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=1000,
                    temperature=0.5,
                    api_key=openapi_key
                )

                # Extract and display the advice
                advice = response.choices[0].message['content'].strip()
                st.subheader("Your Personalized Legal Advice 📜")
                st.write(advice)

            except Exception as e:
                st.error(f"An error occurred: {e}")

        else:
            st.warning("Please fill out all required fields.")


elif selected_option == "Legal Experts":
    st.header("🔍 Connect with Legal Experts")
    st.subheader("Get Expert Legal Assistance")
    st.write("👋 **Welcome!** This feature allows you to connect with qualified legal professionals who can provide personalized advice based on your specific needs.")
    
    # Provide a form for users to connect with legal experts
    with st.form(key='legal_experts_form'):
        st.write("📝 **Fill out the form below to get in touch with a legal expert:**")
        
        name = st.text_input("🧑 Name", placeholder="Enter your full name")
        email = st.text_input("📧 Email", placeholder="Enter your email address")
        phone = st.text_input("📞 Phone Number (optional)", placeholder="Enter your phone number")
        legal_issue = st.text_area("📋 Describe your legal issue", placeholder="Provide a detailed description of your legal issue...")
        
        submit_button = st.form_submit_button("🔗 Connect with Expert")

    if submit_button:
        if name and email and legal_issue:
            try:
                # Prepare the prompt for OpenAI API
                prompt = f"""
                You are a legal assistant helping users connect with legal experts. Use the following information to generate a professional response:

                Name: {name}
                Email: {email}
                Phone: {phone if phone else "None"}
                Legal Issue: {legal_issue}

                Generate a response confirming that the details have been received and that the user will be contacted by a legal expert shortly.
                """

                # Call OpenAI API for generating a confirmation message
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a legal assistant helping users connect with legal experts."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=150,
                    temperature=0.5,
                    api_key=openapi_key
                )

                # Extract and display the confirmation message
                confirmation_message = response.choices[0].message['content'].strip()
                st.success("✅ **Your details have been submitted successfully!**")
                st.write("📧 **Here's what will happen next:**")
                st.write(confirmation_message)

            except Exception as e:
                st.error(f"🚨 **An error occurred:** {e}")

        else:
            st.warning("⚠️ **Please fill out all required fields.**")

elif selected_option == "Compliance Checks":
    st.header("📜 Compliance Checks")
    st.subheader("Ensure Your Documents Meet Legal Standards")
    st.write("🔍 **This feature helps you check whether your documents comply with legal standards.** Please upload your documents for analysis.")

    # File uploader for document submission
    uploaded_file = st.file_uploader("📁 Upload Document", type=["pdf", "docx"], label_visibility="visible", help="Limit 200MB per file • PDF, DOCX")

    if uploaded_file is not None:
        # Display file details
        st.write(f"**Uploaded file:** {uploaded_file.name}")
        st.write("**File type:**", uploaded_file.type)

        # Read the content of the file
        file_content = uploaded_file.read()
        
        # Example placeholder for document processing (actual implementation would vary)
        try:
            # Convert document content to text based on the file type
            if uploaded_file.type == "application/pdf":
                import PyPDF2
                pdf_reader = PyPDF2.PdfFileReader(uploaded_file)
                text = ""
                for page_num in range(pdf_reader.numPages):
                    text += pdf_reader.getPage(page_num).extract_text()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(uploaded_file)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                st.error("Unsupported file type.")
                text = ""

            if text:
                # Call OpenAI API for compliance analysis
                prompt = f"""
                You are an expert in legal compliance. Analyze the following document text and determine whether it complies with standard legal requirements. Provide a summary of the compliance status:

                {text}

                Provide a response summarizing the compliance status of the document.
                """

                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert in legal compliance. Analyze documents for compliance with legal standards."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.5,
                    api_key=openapi_key
                )

                # Extract and display the compliance analysis
                compliance_analysis = response.choices[0].message['content'].strip()
                st.success("✅ **Compliance Analysis Completed!**")
                st.write("**Compliance Summary:**")
                st.write(compliance_analysis)

        except Exception as e:
            st.error(f"🚨 **An error occurred:** {e}")



# Sample data for dropdown options
sample_nodes = ["Case Law", "Statute", "Court", "Judge", "Regulation"]
sample_edges = [
    "Case Law - Statute",
    "Statute - Court",
    "Court - Judge",
    "Judge - Case Law",
    "Regulation - Case Law",
]

if selected_option == "Knowledge Graphs":
    st.header("⚖️ Legal Knowledge Graphs")
    st.subheader("Visualize Legal Concepts and Relationships")

    # User inputs for the knowledge graph
    st.text("Select legal concepts (nodes) and their relationships (edges) for visualization.")

    # Dropdown menu for nodes
    nodes_selected = st.multiselect(
        "Select legal concepts/entities",
        options=sample_nodes,
        default=sample_nodes
    )

    # Dropdown menu for edges
    edges_selected = st.multiselect(
        "Select relationships between entities",
        options=sample_edges,
        default=sample_edges
    )

    # Process selected nodes and edges
    nodes = [node.strip() for node in nodes_selected if node.strip()]
    edges = [tuple(edge.strip().split(' - ')) for edge in edges_selected if ' - ' in edge]

    if nodes and edges:
        # Create graph visualization
        fig = go.Figure()

        # Add edges to the graph
        for edge in edges:
            node1, node2 = edge
            if node1 in nodes and node2 in nodes:
                fig.add_trace(go.Scatter(
                    x=[nodes.index(node1), nodes.index(node2)],
                    y=[nodes.index(node1), nodes.index(node2)],
                    mode='lines+markers',
                    line=dict(width=2, color='blue'),
                    marker=dict(size=10, color='red'),
                    text=[node1, node2],
                    textposition='top center'
                ))

        # Add nodes to the graph
        fig.add_trace(go.Scatter(
            x=list(range(len(nodes))),
            y=list(range(len(nodes))),
            mode='markers+text',
            text=nodes,
            textposition='bottom center',
            marker=dict(size=20, color='orange', line=dict(width=2, color='black')),
            showlegend=False
        ))

        fig.update_layout(
            title="Legal Knowledge Graph",
            title_x=0.5,
            xaxis=dict(showgrid=False, zeroline=False, title="Entities"),
            yaxis=dict(showgrid=False, zeroline=False, title="Entities"),
            showlegend=False,
            plot_bgcolor='white',
            paper_bgcolor='lightgray',
            margin=dict(l=0, r=0, t=50, b=0)
        )

        st.plotly_chart(fig, use_container_width=True)
    else:
        st.warning("Please select legal concepts and relationships to visualize the knowledge graph.")
if selected_option == "Drafting":
    st.header("📄 Automated Legal Document Drafting")
    st.subheader("📝 Assist in drafting standard legal documents.")
    
    # Example fields for document drafting
    st.markdown("### 📋 Enter details for document drafting:")
    
    # Select document type
    doc_type = st.selectbox("📑 Select document type", ["Contract", "Will", "Lease Agreement", "Non-Disclosure Agreement"])
    
    # Input fields for party names and date
    party_one = st.text_input("👤 Party One Name", "John Doe")
    party_two = st.text_input("👥 Party Two Name", "Jane Smith")
    date = st.date_input("📅 Date")
    
    # Sample clauses for different document types
    clauses = {
        "Contract": "🤝 This contract is binding between the parties under the terms agreed upon.",
        "Will": "📜 This will bequeaths the stated assets to the beneficiaries listed.",
        "Lease Agreement": "🏠 This lease agreement outlines the rental terms and obligations of both parties.",
        "Non-Disclosure Agreement": "🔒 This non-disclosure agreement ensures confidentiality of the shared information."
    }
    
    if st.button("🖋️ Draft Document"):
        if doc_type and party_one and party_two and date:
            # Draft document based on user inputs and selected document type
            document = f"""
            {doc_type} 📄
            
            This {doc_type.lower()} is made on {date} between {party_one} and {party_two}.
            
            {clauses[doc_type]}
            
            ⚖️ Terms and conditions apply.
            
            🖋️ Signed,
            {party_one}
            {party_two}
            """
            st.text_area("📃 Drafted Document", document, height=300)
        else:
            st.error("🚨 Please fill in all fields to draft the document.")



if selected_option == "Precedents":
    st.header("📚 Legal Precedent Finder")
    st.subheader("🔍 Find relevant case laws and precedents")

    # Example fields for precedent finding
    st.markdown("### 📝 Enter details to find legal precedents:")

    # Input fields for case details
    case_topic = st.text_input("⚖️ Case Topic", "Contract Law")
    jurisdiction = st.text_input("🌍 Jurisdiction", "United States")
    year = st.number_input("📅 Year", min_value=1800, max_value=2024, step=1, value=2024)

    # Example case precedents for demonstration purposes
    sample_precedents = [
        {"case": "Brown v. Board of Education", "year": 1954, "jurisdiction": "United States", "summary": "Landmark case declaring state laws establishing separate public schools for black and white students to be unconstitutional."},
        {"case": "Roe v. Wade", "year": 1973, "jurisdiction": "United States", "summary": "Landmark decision by the US Supreme Court that ruled that the Constitution protects a pregnant woman's liberty to choose to have an abortion."},
        {"case": "Obergefell v. Hodges", "year": 2015, "jurisdiction": "United States", "summary": "Landmark civil rights case in which the Supreme Court ruled that the fundamental right to marry is guaranteed to same-sex couples."}
    ]

    # Filter precedents based on user inputs
    if st.button("🔍 Find Precedents"):
        if case_topic and jurisdiction and year:
            filtered_precedents = [precedent for precedent in sample_precedents if 
                                   precedent["jurisdiction"].lower() == jurisdiction.lower() and
                                   precedent["year"] <= year]

            if filtered_precedents:
                for precedent in filtered_precedents:
                    st.markdown(f"""
                    **📜 Case:** {precedent['case']}
                    **📅 Year:** {precedent['year']}
                    **🌍 Jurisdiction:** {precedent['jurisdiction']}
                    **📝 Summary:** {precedent['summary']}
                    """)
            else:
                st.warning("⚠️ No precedents found for the given criteria.")
        else:
            st.error("🚨 Please fill in all fields to find legal precedents.")



# Custom CSS for LawGPT
st.markdown("""
    <style>
      /* Global Styles */
.stApp {
    background: linear-gradient(to right, #e0f7fa, #80deea); /* Main app gradient */
    font-family: 'Arial', sans-serif;
}

/* Sidebar Styles */
.sidebar {
    background: linear-gradient(to right, #004d40, #00796b); /* Sidebar gradient */
    color: #ffffff; /* Text color for readability */
    border-right: 2px solid #004d40; /* Border color matching the sidebar gradient */
    padding: 20px; /* Padding for spacing */
    height: 100vh; /* Full viewport height */
    overflow-y: auto; /* Scroll if content overflows */
    position: fixed; /* Fixed position for the sidebar */
    width: 250px; /* Fixed width for the sidebar */
    box-shadow: 4px 0 8px rgba(0, 0, 0, 0.1); /* Subtle shadow effect */
}

/* Sidebar Links */
.sidebar a {
    color: #ffffff; /* Link color for readability */
    text-decoration: none; /* Remove underline from links */
    display: block; /* Display links as block elements */
    padding: 10px 15px; /* Padding for clickable areas */
    border-radius: 5px; /* Rounded corners for links */
    transition: background-color 0.3s, color 0.3s; /* Smooth transition for hover effects */
}

.sidebar a:hover {
    background-color: #00796b; /* Darker background on hover */
    color: #ffffff; /* Keep text color white on hover */
}

.sidebar .active {
    background-color: #004d40; /* Active link background */
    color: #ffffff; /* Active link text color */
}

/* Sidebar Section Titles */
.sidebar .title {
    font-size: 1.2em; /* Larger font size for section titles */
    margin-bottom: 20px; /* Spacing below title */
    color: #ffffff; /* Title text color */
}

/* Radio Buttons */
.stRadio {
    display: flex;
    flex-direction: column;
    gap: 10px; /* Space between radio buttons */
}

.stRadio label {
    color: #004d40; /* Color for radio button labels */
    font-size: 1em; /* Font size for labels */
}

.stRadio input[type="radio"] {
    accent-color: #004d40; /* Color for radio buttons */
}

/* Text Inputs */
.stTextInput input {
    background-color: #ffffff; /* Background color for input fields */
    color: #004d40; /* Text color for input fields */
    border: 1px solid #004d40; /* Border color matching the main theme */
    border-radius: 5px; /* Rounded corners for input fields */
    padding: 10px; /* Padding inside input fields */
}

/* Buttons */
.stButton {
    background-color: #004d40; /* Button background color */
    color: black; /* Button text color */
    border: none; /* Remove default border */
    border-radius: 5px; /* Rounded corners for buttons */
    padding: 10px 20px; /* Padding inside buttons */
    cursor: pointer; /* Pointer cursor on hover */
    # transition: background-color 0.3s; /* Smooth transition for hover effects */
}

.stButton:hover {
    background-color: #00796b; /* Darker background on hover */
}

            
    </style>
""", unsafe_allow_html=True)
